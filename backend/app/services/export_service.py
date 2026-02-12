"""病历导出服务 — PDF / Word / Excel"""

import io
import uuid
from datetime import datetime

from sqlalchemy.ext.asyncio import AsyncSession

from app.models.medical_record import MedicalRecord
from app.services.record_service import get_record, list_records
from app.models.user import User


# ---- 通用辅助 ----

SEVERITY_MAP = {
    "mild": "轻度",
    "moderate": "中度",
    "severe": "重度",
    "critical": "危急",
}

FIELD_LABELS = {
    "record_no": "病历编号",
    "visit_date": "就诊日期",
    "poultry_type": "禽类类型",
    "breed": "品种",
    "age_days": "日龄",
    "affected_count": "发病数",
    "total_flock": "总群数",
    "onset_date": "发病日期",
    "primary_diagnosis": "主要诊断",
    "severity": "严重程度",
    "icd_code": "ICD编码",
    "status": "状态",
    "current_version": "版本",
}


def _flatten_record(record: MedicalRecord) -> dict:
    """将病历对象扁平化为字典"""
    return {
        "record_no": record.record_no,
        "visit_date": str(record.visit_date),
        "poultry_type": record.poultry_type,
        "breed": record.breed or "",
        "age_days": str(record.age_days) if record.age_days else "",
        "affected_count": str(record.affected_count) if record.affected_count else "",
        "total_flock": str(record.total_flock) if record.total_flock else "",
        "onset_date": str(record.onset_date) if record.onset_date else "",
        "primary_diagnosis": record.primary_diagnosis or "",
        "severity": SEVERITY_MAP.get(record.severity, record.severity or ""),
        "icd_code": record.icd_code or "",
        "status": record.status,
        "current_version": record.current_version,
    }


def _extract_sections(record_json: dict) -> list[tuple[str, str]]:
    """从 record_json 提取各节内容"""
    sections = []

    if "basic_info" in record_json:
        info = record_json["basic_info"]
        text = "\n".join(f"{k}: {v}" for k, v in info.items())
        sections.append(("基本信息", text))

    if "symptoms" in record_json:
        symptoms = record_json["symptoms"]
        if isinstance(symptoms, list):
            text = "\n".join(f"• {s}" for s in symptoms)
        elif isinstance(symptoms, dict):
            text = "\n".join(f"• {k}: {v}" for k, v in symptoms.items())
        else:
            text = str(symptoms)
        sections.append(("症状", text))

    if "primary_diagnosis" in record_json:
        diag = record_json["primary_diagnosis"]
        sev = record_json.get("severity", "")
        text = f"诊断: {diag}"
        if sev:
            text += f"\n严重程度: {SEVERITY_MAP.get(sev, sev)}"
        sections.append(("诊断", text))

    if "treatment" in record_json:
        treatment = record_json["treatment"]
        if isinstance(treatment, dict):
            text = "\n".join(f"{k}: {v}" for k, v in treatment.items())
        elif isinstance(treatment, list):
            text = "\n".join(f"• {t}" for t in treatment)
        else:
            text = str(treatment)
        sections.append(("治疗方案", text))

    if "vaccination_history" in record_json:
        sections.append(("免疫史", str(record_json["vaccination_history"])))

    if "environment" in record_json:
        sections.append(("环境条件", str(record_json["environment"])))

    if "notes" in record_json:
        sections.append(("备注", str(record_json["notes"])))

    return sections


# ---- PDF 导出 ----

async def export_record_pdf(db: AsyncSession, record_id: uuid.UUID) -> bytes:
    """导出单个病历为 PDF"""
    record = await get_record(db, record_id)
    flat = _flatten_record(record)
    sections = _extract_sections(record.record_json)

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=20*mm, bottomMargin=20*mm)
    styles = getSampleStyleSheet()

    # 尝试注册中文字体（如果可用）
    try:
        import os
        font_paths = [
            "C:/Windows/Fonts/simhei.ttf",
            "C:/Windows/Fonts/simsun.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/System/Library/Fonts/PingFang.ttc",
        ]
        for fp in font_paths:
            if os.path.exists(fp):
                pdfmetrics.registerFont(TTFont("ChineseFont", fp))
                styles.add(ParagraphStyle("CN", fontName="ChineseFont", fontSize=10, leading=14))
                styles.add(ParagraphStyle("CNTitle", fontName="ChineseFont", fontSize=16, leading=20, spaceAfter=10))
                styles.add(ParagraphStyle("CNH2", fontName="ChineseFont", fontSize=12, leading=16, spaceAfter=6, spaceBefore=8))
                break
    except Exception:
        pass

    cn = styles.get("CN", styles["Normal"])
    cn_title = styles.get("CNTitle", styles["Title"])
    cn_h2 = styles.get("CNH2", styles["Heading2"])

    elements = []
    elements.append(Paragraph(f"禽类病历报告", cn_title))
    elements.append(Spacer(1, 5*mm))

    # 基本信息表格
    table_data = []
    for key in ("record_no", "visit_date", "poultry_type", "breed", "age_days",
                "affected_count", "total_flock", "primary_diagnosis", "severity", "status"):
        label = FIELD_LABELS.get(key, key)
        value = flat.get(key, "")
        table_data.append([label, value])

    if table_data:
        t = Table(table_data, colWidths=[80*mm, 90*mm])
        t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (0, -1), colors.Color(0.95, 0.95, 0.95)),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ]))
        elements.append(t)

    # 详细信息各节
    for title, content in sections:
        elements.append(Spacer(1, 4*mm))
        elements.append(Paragraph(title, cn_h2))
        for line in content.split("\n"):
            if line.strip():
                elements.append(Paragraph(line, cn))

    elements.append(Spacer(1, 10*mm))
    elements.append(Paragraph(
        f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        cn,
    ))

    doc.build(elements)
    return buf.getvalue()


# ---- Word 导出 ----

async def export_record_word(db: AsyncSession, record_id: uuid.UUID) -> bytes:
    """导出单个病历为 Word 文档"""
    record = await get_record(db, record_id)
    flat = _flatten_record(record)
    sections = _extract_sections(record.record_json)

    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 标题
    title = doc.add_heading("禽类病历报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息表格
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key in ("record_no", "visit_date", "poultry_type", "breed", "age_days",
                "affected_count", "total_flock", "onset_date",
                "primary_diagnosis", "severity", "icd_code", "status", "current_version"):
        label = FIELD_LABELS.get(key, key)
        value = flat.get(key, "")
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value

    # 详细信息
    for title_text, content in sections:
        doc.add_heading(title_text, level=2)
        doc.add_paragraph(content)

    # 导出时间
    doc.add_paragraph("")
    footer = doc.add_paragraph(
        f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    footer.runs[0].font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---- Excel 批量导出 ----

async def export_records_excel(
    db: AsyncSession, user: User,
    record_ids: list[uuid.UUID] | None = None,
) -> bytes:
    """批量导出病历为 Excel"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "病历列表"

    # 表头
    headers = ["病历编号", "就诊日期", "禽类", "品种", "日龄", "发病数",
               "总群数", "发病日期", "主要诊断", "严重程度", "状态", "版本"]
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    # 获取病历数据
    if record_ids:
        from sqlalchemy import select
        from app.models.medical_record import MedicalRecord
        result = await db.execute(
            select(MedicalRecord).where(
                MedicalRecord.id.in_(record_ids),
                MedicalRecord.status != "deleted",
            )
        )
        records = list(result.scalars().all())
    else:
        records, _ = await list_records(db, user, page=1, page_size=1000)

    # 数据行
    for row_idx, record in enumerate(records, 2):
        flat = _flatten_record(record)
        values = [
            flat["record_no"], flat["visit_date"], flat["poultry_type"],
            flat["breed"], flat["age_days"], flat["affected_count"],
            flat["total_flock"], flat["onset_date"], flat["primary_diagnosis"],
            flat["severity"], flat["status"], flat["current_version"],
        ]
        for col, value in enumerate(values, 1):
            ws.cell(row=row_idx, column=col, value=value)

    # 自动列宽
    for col_idx in range(1, len(headers) + 1):
        max_len = max(
            len(str(ws.cell(row=r, column=col_idx).value or ""))
            for r in range(1, ws.max_row + 1)
        )
        ws.column_dimensions[chr(64 + col_idx)].width = min(max_len + 4, 30)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
